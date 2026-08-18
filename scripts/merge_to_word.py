#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: Apache-2.0
# Copyright 2026 newdeme
"""
merge_to_word.py —— 书稿合并 Word 生成器（monograph-writing 技能组件）
========================================================================
用途：把剥离版书稿（配置 stripped_dir，缺省 04_剥离版书稿/）按《专著目录.md》的
结构合并为**一个 Word 文档**：
  - 已完成章节填入正文；未完成章节保留完整标题骨架并标注"待并入"，
    使目录与专著目录完全一致；
  - 标题层级：章 = Heading 1（居中、黑体、章前分页）、节 = Heading 2、小节 = Heading 3；
    Word 目录（TOC 域）基于 Heading 1~3 生成；
  - 每个小节单元 = 小节标题＋正文＋该小节参考文献；节级总结以【本节小结】
    （非标题样式、不入目录）附于该节末尾。

版式：A4；正文宋体小四（12pt）1.5 倍行距、首行缩进 2 字符；标题黑体加粗黑色；
参考文献五号（10.5pt）悬挂缩进；表格 Table Grid 带框线、表头加粗；页脚居中页码。

依赖：python-docx（只需这一步安装：pip3 install python-docx）

用法（先跑 generate_stripped_version.py 生成剥离版，再跑本脚本）：
    python3 merge_to_word.py --root .
"""
import argparse
import json
import os
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[缺依赖] 生成 Word 需要安装 python-docx，请在终端执行：")
    print("    pip3 install python-docx")
    print("（装好后重新运行本脚本即可；不确定怎么装就把这句发给 AI 助手代跑。）")
    sys.exit(1)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CH_RE = re.compile(r"^## (第.+?章 .+)$")
SEC_RE = re.compile(r"^\*\*(\d+\.\d+) (.+?)\*\*$")
ITEM_RE = re.compile(r"^- (\d+\.\d+\.\d+) (.+)$")
CAPTION_RE = re.compile(r"^(\*\*)?[表图]\s*\d+-\d+")

MISSING_NOTE = "（本章正文尚未完成并入，本版暂缺。）"
GRAY = RGBColor(0x80, 0x80, 0x80)

FUZZY = []  # "目录标题含文件系统非法字符、按编号前缀回退匹配"的记录


def load_config(root: Path):
    cfg = {"book_title": "未命名书稿", "catalog_file": "00_管理文件/专著目录.md",
           "stripped_dir": "04_剥离版书稿", "author": ""}
    cfile = root / "00_管理文件" / "书稿配置.json"
    if cfile.is_file():
        user = json.loads(cfile.read_text(encoding="utf-8"))
        cfg.update({k: user[k] for k in cfg if k in user and user[k]})
    return cfg


def parse_catalog(path):
    """目录 → [(章标题, [(节号, 节标题, [(小节号, 小节标题)...])...])...]"""
    chapters = []
    cur_ch, cur_sec = None, None
    for line in open(path, encoding="utf-8").read().splitlines():
        line = line.rstrip()
        m = CH_RE.match(line)
        if m:
            cur_ch = {"title": m.group(1), "sections": []}
            chapters.append(cur_ch)
            cur_sec = None
            continue
        m = SEC_RE.match(line)
        if m and cur_ch is not None:
            cur_sec = {"num": m.group(1), "title": m.group(2), "items": []}
            cur_ch["sections"].append(cur_sec)
            continue
        m = ITEM_RE.match(line)
        if m and cur_sec is not None:
            cur_sec["items"].append((m.group(1), m.group(2)))
    return chapters


def locate(src, ch_title, sec, item=None, kind="item"):
    """在剥离版镜像树中定位文件；目录标题含'/'等非法字符时按编号前缀回退匹配。"""
    cd = os.path.join(src, ch_title)
    if kind == "item":
        sd = os.path.join(cd, f"{sec['num']} {sec['title']}")
        exact = os.path.join(sd, f"{item[0]} {item[1]}.md")
        if os.path.isfile(exact):
            return exact
        cands = [f for f in os.listdir(sd)
                 if f.startswith(item[0] + " ") and f.endswith(".md")]
        if len(cands) == 1:
            FUZZY.append(f"{item[0]}（目录：{item[1]} ｜ 文件：{cands[0][:-3]}）")
            return os.path.join(sd, cands[0])
        raise FileNotFoundError(f"小节 {item[0]} 无法唯一定位（候选 {cands}）：{sd}")
    if kind == "sec_summary":
        sd = os.path.join(cd, f"{sec['num']} {sec['title']}")
        return os.path.join(
            sd, f"{ch_title.split()[0]} {sec['num']}节 {sec['title']} 章节总结.md")
    if kind == "ch_summary":
        stem = os.path.join(cd, f"{sec['num']} 小结")
        return stem + ".md" if os.path.isfile(stem + ".md") \
            else os.path.join(stem, f"{sec['num']} 小结.md")


def read_unit(path):
    """读取剥离文件 → (正文行列表, 参考文献行列表)；H1 行丢弃。"""
    lines = open(path, encoding="utf-8").read().splitlines()
    try:
        ri = next(i for i, l in enumerate(lines) if l.strip() == REF_HEADING)
    except StopIteration:
        raise ValueError(f"未找到参考文献分隔：{path}")
    body = [l for l in lines[1:ri] if l.strip() and l.strip() != "---"]
    refs = [l for l in lines[ri + 1:] if l.strip() and l.strip() != "---"]
    return body, refs


REF_HEADING = "## 参考文献"


# ---------------------------------------------------------------- 样式
def set_eastasia(style, ea):
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastasia"), ea)


def setup_styles(doc):
    st = doc.styles["Normal"]
    st.font.size = Pt(12)
    set_eastasia(st, "宋体")
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(24)
    pf.space_after = Pt(0)

    for name, size, center in (("Heading 1", 16, True), ("Heading 2", 14, False),
                               ("Heading 3", 12, False)):
        s = doc.styles[name]
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        set_eastasia(s, "黑体")
        s.paragraph_format.space_before = Pt(12 if size > 12 else 6)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.first_line_indent = Pt(0)
        s.paragraph_format.keep_with_next = True
        if center:
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            s.paragraph_format.page_break_before = True
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(fp, "PAGE", "1")


def _field(paragraph, instr, placeholder):
    r1 = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._r.append(f1)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "; r2._r.append(it)
    r3 = paragraph.add_run()
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate"); r3._r.append(f2)
    paragraph.add_run(placeholder)
    r4 = paragraph.add_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end"); r4._r.append(f3)


def add_runs(p, text, base_bold=False):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()]); r.bold = base_bold or None
        r = p.add_run(m.group(1)); r.bold = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); r.bold = base_bold or None
    if not p.runs:
        p.add_run(text)


def para(doc, text, style=None, center=False, no_indent=False, size=None,
         color=None, space_before=None, base_bold=False):
    p = doc.add_paragraph(style=style)
    add_runs(p, text, base_bold=base_bold)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if no_indent:
        p.paragraph_format.first_line_indent = Pt(0)
    if size is not None:
        for r in p.runs:
            r.font.size = Pt(size)
    if color is not None:
        for r in p.runs:
            r.font.color.rgb = color
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


# ---------------------------------------------------------------- 内容写入
def write_body(doc, body_lines):
    i = 0
    while i < len(body_lines):
        raw = body_lines[i]
        was_quote = raw.lstrip().startswith(">")
        ln = raw.lstrip()[1:].lstrip() if was_quote else raw
        s = ln.strip()
        if s.startswith("|"):
            blk = []
            while i < len(body_lines):
                r2 = body_lines[i]
                r2 = r2.lstrip()[1:].lstrip() if r2.lstrip().startswith(">") else r2
                if r2.strip().startswith("|"):
                    blk.append(r2); i += 1
                else:
                    break
            add_table(doc, blk)
            continue
        if s.startswith("$$") and s.endswith("$$"):
            para(doc, s[2:-2].strip(), center=True, no_indent=True)
        elif s.startswith("**") and s.count("**") == 2 and CAPTION_RE.match(s):
            para(doc, s, center=True, no_indent=True, size=10.5)
        elif was_quote:
            para(doc, s, no_indent=True, size=10.5, color=GRAY)
        elif s.startswith("### "):
            para(doc, s[4:], no_indent=True, base_bold=True)
        else:
            para(doc, ln)
        i += 1


def add_table(doc, tbl_lines):
    rows = []
    for ln in tbl_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.match(r"^:?-{2,}:?$", c) for c in cells if c):
            continue
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        for j in range(ncols):
            cell = t.cell(i, j)
            cp = cell.paragraphs[0]
            cp.paragraph_format.first_line_indent = Pt(0)
            cp.paragraph_format.line_spacing = 1.0
            txt = r[j] if j < len(r) else ""
            add_runs(cp, txt, base_bold=(i == 0))
            for run in cp.runs:
                run.font.size = Pt(10.5)


def write_refs(doc, refs):
    p = para(doc, "参考文献", no_indent=True, space_before=10)
    for r in p.runs:
        r.bold = True
        r.font.size = Pt(10.5)
    for ln in refs:
        p = para(doc, ln, no_indent=True)
        pf = p.paragraph_format
        pf.left_indent = Pt(21); pf.first_line_indent = Pt(-21)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)


def write_unit(doc, path):
    body, refs = read_unit(path)
    write_body(doc, body)
    write_refs(doc, refs)


# ---------------------------------------------------------------- 主流程
def build(root):
    cfg = load_config(root)
    src = root / cfg["stripped_dir"]
    catalog = root / cfg["catalog_file"]
    if not catalog.is_file():
        print(f"[目录缺失] {catalog}（专著目录是合并结构的依据，必须有）")
        sys.exit(1)
    if not src.is_dir():
        print(f"[剥离版缺失] {src}——请先运行 generate_stripped_version.py")
        sys.exit(1)
    out = src / f"{cfg['book_title']}（合并稿）.docx"

    chapters = parse_catalog(catalog)
    doc = Document()
    setup_styles(doc)
    doc.core_properties.title = cfg["book_title"]
    if cfg.get("author"):
        doc.core_properties.author = cfg["author"]

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(f"《{cfg['book_title']}》")
    r.font.size = Pt(22); r.bold = True
    r.font.name = "Times New Roman"
    r._r.rPr.rFonts.set(qn("w:eastasia"), "黑体")
    para(doc, f"（合并稿｜章节结构按专著目录｜{date.today()} 生成）",
         center=True, no_indent=True, size=10.5, color=GRAY)

    hp = doc.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("目  录"); hr.font.size = Pt(16); hr.bold = True
    hr.font.name = "Times New Roman"
    hr._r.rPr.rFonts.set(qn("w:eastasia"), "黑体")
    hp.paragraph_format.page_break_before = True
    toc_p = doc.add_paragraph()
    _field(toc_p, 'TOC \\o "1-3" \\h \\z \\u',
           "（目录域：在 Word 中全选 Ctrl+A 后按 F9 更新，即可生成全书目录及页码）")

    stats = {"ch": 0, "items": 0, "sec_sum": 0, "missing": [], "tables": 0}
    for ch in chapters:
        ch_title = ch["title"]
        doc.add_heading(ch_title, level=1)
        cd = os.path.join(src, ch_title)
        have = os.path.isdir(cd)
        if not have:
            stats["missing"].append(ch_title)
            para(doc, MISSING_NOTE, no_indent=True, color=GRAY)
        stats["ch"] += 1
        for sec in ch["sections"]:
            doc.add_heading(f"{sec['num']} {sec['title']}", level=2)
            for item in sec["items"]:
                doc.add_heading(f"{item[0]} {item[1]}", level=3)
                if have:
                    p = locate(str(src), ch_title, sec, item, "item")
                    if not os.path.isfile(p):
                        print(f"[提示] 缺少小节文件（跳过）：{p}")
                        para(doc, "（本小节尚未完成，待并入。）", no_indent=True, color=GRAY)
                        continue
                    write_unit(doc, p)
                    stats["items"] += 1
                else:
                    para(doc, "（本小节尚未完成，待并入。）", no_indent=True, color=GRAY)
            is_end = sec["title"] == "小结"
            if have and is_end:
                p = locate(str(src), ch_title, sec, kind="ch_summary")
                if os.path.isfile(p):
                    write_unit(doc, p)
            elif have and sec["items"]:
                p = locate(str(src), ch_title, sec, kind="sec_summary")
                if os.path.isfile(p):
                    lead = para(doc, "【本节小结】", no_indent=True, space_before=10)
                    for r0 in lead.runs:
                        r0.bold = True
                    write_unit(doc, p)
                    stats["sec_sum"] += 1
    stats["tables"] = len(doc.tables)
    stats["fuzzy"] = list(FUZZY)
    doc.save(out)
    return out, stats


def main():
    ap = argparse.ArgumentParser(description="合并 Word 稿生成（详细说明见文件头注释）")
    ap.add_argument("--root", default=".", help="项目根目录")
    args = ap.parse_args()
    root = Path(args.root).expanduser().resolve()
    out, stats = build(root)
    print(f"生成完成 → {out}")
    print(f"章: {stats['ch']}（待并入: {len(stats['missing'])}）")
    for m0 in stats["missing"]:
        print(f"  - {m0}")
    print(f"已并入小节: {stats['items']}｜节级小结: {stats['sec_sum']}｜表格: {stats['tables']}")
    if stats.get("fuzzy"):
        print(f"编号回退匹配: {len(stats['fuzzy'])} 条（目录标题含文件系统非法字符，Word 标题仍用目录原文）")
        for f in stats["fuzzy"]:
            print(f"  - {f}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
